import os
import pandas as pd
from docx import Document
from docx.shared import RGBColor

# Load the data from the data source (Excel file) with column names
file_path = 'DataSource.xlsx'
if not os.path.exists(file_path):
    print('DataSource not found. Please make sure it is correctly named: "DataSource", is an Excel file, and exists in the same folder as the python script.\n'
      'Alternatively,if you have customized the file name, format(.csv) or location, make the appropriate changes in the script.')
    input('Please press Enter to close the program, run it again after making sure the file is here.')
    exit()

source_data = pd.read_excel(file_path, engine='openpyxl', sheet_name='List1', names=['English', 'Farsi'])


# Convert any non-string values to strings in the source data
source_data['English'] = source_data['English'].astype(str)
source_data['Farsi'] = source_data['Farsi'].astype(str)

# Convert English sentences in source_data to lowercase
source_data['English'] = source_data['English'].str.lower()

# Load the content from the given file
doc_path = input("Please paste the word file's full path(file extension included):\n")
while not os.path.exists(doc_path) or doc_path[-4:] != 'docx':
    doc_path = input('file not found, please make sure the full correct path is provided and the document is a docx file. try again:\n')
doc = Document(doc_path)

# Modify the find_replace function to change the color
def find_replace(text, source_data):

    translated_text = Document()  # Create a new Document object
    paragraph_text = ""  # Accumulate text within a paragraph

    for paragraph in text.paragraphs:
        new_paragraph = translated_text.add_paragraph()
        paragraph_text = ""  # Reset the accumulated text for each paragraph

        for run in paragraph.runs:
            run_text = run.text
            paragraph_text += run_text

            for index, row in source_data.iterrows():
                english_sentence = row['English']
                farsi_translation = row['Farsi']

                # Convert the English sentence to lowercase for a case-insensitive comparison
                if english_sentence.lower() in paragraph_text.lower():
                    # Match found, add both English and Farsi text
                    new_run = new_paragraph.add_run('\n' + paragraph_text + '\n' + paragraph_text.lower().replace(english_sentence.lower(), farsi_translation) + '\n')
                    paragraph_text = ""  # Reset accumulated text
                    break  # Exit the loop after the first match

        if paragraph_text:  # If there's any remaining text in the paragraph
            # No match found, add English sentence with red color
            new_run = new_paragraph.add_run('Match not found:\n' + paragraph_text + '\n')
            font = new_run.font
            font.color.rgb = RGBColor(255, 0, 0)  # Red color

    return translated_text

# Process the content by finding and replacing character combinations
translated_content = find_replace(doc, source_data)

#function to use when handling permission exception:
def is_file_open(file_path):
    try:
        with open(file_path, 'r'):
            return False  # The file is not open
    except PermissionError:
        return True  # The file is open by another process
    

# Save the output document with the translated content
output_path = "Translated_" + doc_path

if os.path.exists(output_path):
    user_choice = input(f"'{output_path}' already exists. Do you want to replace it? (Y/N): ").strip().lower()
    if user_choice == 'y':
            # User chose to replace the file, proceed without removing it
        pass
    else:
            # User chose not to replace the file, generate a new file name
        num = 1
        while os.path.exists(f"Translated_{num}_{doc_path}"):
            num += 1
        output_path = f"Translated_{num}_{doc_path}"
        print(f"Saving the translated document as '{output_path}'.")


try:
    translated_content.save(output_path)  # Save the Document object
except:
        print("An error occurred: the file is probably already opened by another program. Make sure it's closed and try again.")
        input('press Enter to exit')
        exit()

print(f"Translated document saved as '{output_path}'.")
input('Press Enter to exit:')
